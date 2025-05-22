import tkinter as tk
from tkinter import filedialog, font, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

class DocxEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Editor de Documentos DOCX")
        self.root.geometry("800x600")
        self.current_file = None
        self.setup_ui()
        
    def setup_ui(self):
        # Barra de herramientas
        toolbar = tk.Frame(self.root, bd=1, relief=tk.RAISED)
        toolbar.pack(side=tk.TOP, fill=tk.X)
        
        # Botones de archivo
        tk.Button(toolbar, text="Abrir", command=self.open_file).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(toolbar, text="Guardar", command=self.save_file).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(toolbar, text="Guardar como", command=self.save_as_file).pack(side=tk.LEFT, padx=2, pady=2)
        
        # Separador
        tk.Label(toolbar, text="|").pack(side=tk.LEFT, padx=5)
        
        # Botones de formato
        self.bold_btn = tk.Button(toolbar, text="Negrita", command=lambda: self.toggle_format("bold"))
        self.bold_btn.pack(side=tk.LEFT, padx=2, pady=2)
        
        self.italic_btn = tk.Button(toolbar, text="Cursiva", command=lambda: self.toggle_format("italic"))
        self.italic_btn.pack(side=tk.LEFT, padx=2, pady=2)
        
        self.underline_btn = tk.Button(toolbar, text="Subrayado", command=lambda: self.toggle_format("underline"))
        self.underline_btn.pack(side=tk.LEFT, padx=2, pady=2)
        
        # Selector de fuente
        self.font_family = tk.StringVar()
        self.font_family.set("Arial")
        self.font_family.trace("w", self.change_font)
        font_families = list(font.families())
        font_families.sort()
        tk.OptionMenu(toolbar, self.font_family, *font_families[:10]).pack(side=tk.LEFT, padx=2, pady=2)
        
        # Selector de tamaño
        self.font_size = tk.IntVar()
        self.font_size.set(12)
        self.font_size.trace("w", self.change_font)
        tk.OptionMenu(toolbar, self.font_size, *range(8, 73, 2)).pack(side=tk.LEFT, padx=2, pady=2)
        
        # Alineación
        tk.Button(toolbar, text="Izquierda", command=lambda: self.set_alignment("left")).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(toolbar, text="Centro", command=lambda: self.set_alignment("center")).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(toolbar, text="Derecha", command=lambda: self.set_alignment("right")).pack(side=tk.LEFT, padx=2, pady=2)
        
        # Área de texto
        self.text_area = tk.Text(self.root, wrap=tk.WORD, undo=True)
        self.text_area.pack(expand=True, fill=tk.BOTH)
        
        # Barra de estado
        self.status_bar = tk.Label(self.root, text="Listo", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Configurar tags para formato
        self.text_area.tag_configure("bold", font=("Arial", 12, "bold"))
        self.text_area.tag_configure("italic", font=("Arial", 12, "italic"))
        self.text_area.tag_configure("underline", font=("Arial", 12, "underline"))
        
        # Menú
        self.menu_bar = tk.Menu(self.root)
        self.file_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.file_menu.add_command(label="Abrir", command=self.open_file)
        self.file_menu.add_command(label="Guardar", command=self.save_file)
        self.file_menu.add_command(label="Guardar como", command=self.save_as_file)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="Salir", command=self.root.quit)
        self.menu_bar.add_cascade(label="Archivo", menu=self.file_menu)
        
        self.format_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.format_menu.add_command(label="Negrita", command=lambda: self.toggle_format("bold"), accelerator="Ctrl+B")
        self.format_menu.add_command(label="Cursiva", command=lambda: self.toggle_format("italic"), accelerator="Ctrl+I")
        self.format_menu.add_command(label="Subrayado", command=lambda: self.toggle_format("underline"), accelerator="Ctrl+U")
        self.menu_bar.add_cascade(label="Formato", menu=self.format_menu)
        
        self.root.config(menu=self.menu_bar)
        
        # Atajos de teclado
        self.root.bind("<Control-b>", lambda e: self.toggle_format("bold"))
        self.root.bind("<Control-i>", lambda e: self.toggle_format("italic"))
        self.root.bind("<Control-u>", lambda e: self.toggle_format("underline"))
        
    def open_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos Word", "*.docx"), ("Todos los archivos", "*.*")]
        )
        
        if file_path:
            try:
                doc = Document(file_path)
                self.text_area.delete(1.0, tk.END)
                
                for para in doc.paragraphs:
                    self.text_area.insert(tk.END, para.text + "\n")
                
                self.current_file = file_path
                self.status_bar.config(text=f"Archivo abierto: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo abrir el archivo:\n{str(e)}")
    
    def save_file(self):
        if not self.current_file:
            self.save_as_file()
            return
            
        try:
            doc = Document()
            text = self.text_area.get(1.0, tk.END)
            
            # Guardar con formato básico (el formato completo requeriría más lógica)
            for line in text.split("\n"):
                para = doc.add_paragraph(line)
                # Aquí podríamos añadir lógica para aplicar formatos específicos
                
            doc.save(self.current_file)
            self.status_bar.config(text=f"Archivo guardado: {os.path.basename(self.current_file)}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar el archivo:\n{str(e)}")
    
    def save_as_file(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx"), ("Todos los archivos", "*.*")]
        )
        
        if file_path:
            self.current_file = file_path
            self.save_file()
    
    def toggle_format(self, format_type):
        try:
            current_tags = self.text_area.tag_names(tk.SEL_FIRST)
            
            if format_type in current_tags:
                self.text_area.tag_remove(format_type, tk.SEL_FIRST, tk.SEL_LAST)
            else:
                self.text_area.tag_add(format_type, tk.SEL_FIRST, tk.SEL_LAST)
                
            # Actualizar botones de la barra de herramientas
            for btn, tag in [(self.bold_btn, "bold"), 
                            (self.italic_btn, "italic"), 
                            (self.underline_btn, "underline")]:
                if tag in self.text_area.tag_names("sel.first"):
                    btn.config(relief=tk.SUNKEN)
                else:
                    btn.config(relief=tk.RAISED)
                    
        except tk.TclError:
            pass  # No hay texto seleccionado
    
    def change_font(self, *args):
        try:
            # Obtener el formato actual del texto seleccionado
            current_font = font.Font(font=self.text_area.cget("font"))
            new_font = (self.font_family.get(), self.font_size.get())
            
            # Configurar el tag con la nueva fuente
            self.text_area.tag_configure("custom_font", font=new_font)
            
            # Aplicar el tag al texto seleccionado
            self.text_area.tag_add("custom_font", tk.SEL_FIRST, tk.SEL_LAST)
            
        except tk.TclError:
            pass  # No hay texto seleccionado
    
    def set_alignment(self, align):
        try:
            # Esta función es simplificada - en un editor real necesitarías más lógica
            # para manejar párrafos en DOCX
            if align == "left":
                self.text_area.tag_configure("align_left", justify=tk.LEFT)
                self.text_area.tag_add("align_left", "sel.first", "sel.last")
            elif align == "center":
                self.text_area.tag_configure("align_center", justify=tk.CENTER)
                self.text_area.tag_add("align_center", "sel.first", "sel.last")
            elif align == "right":
                self.text_area.tag_configure("align_right", justify=tk.RIGHT)
                self.text_area.tag_add("align_right", "sel.first", "sel.last")
                
        except tk.TclError:
            pass  # No hay texto seleccionado

if __name__ == "__main__":
    root = tk.Tk()
    editor = DocxEditor(root)
    root.mainloop()